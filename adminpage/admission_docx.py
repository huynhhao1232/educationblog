import io
import os
import re

from django.conf import settings
from docx import Document

ADMISSION_DOCX_TEMPLATE = os.path.join(
    settings.BASE_DIR,
    'tuyển sinh',
    'ADMISSION_FORM_TEMPLATE.docx',
)


def _fmt_date(value):
    if not value:
        return '.../.../......'
    return value.strftime('%d/%m/%Y')


def _fmt_sign_date(value):
    if not value:
        return 'ngày ... tháng ... năm 2026'
    return f'ngày {value.day:02d} tháng {value.month:02d} năm {value.year}'


def _graduation_year_label(value):
    if value == 'current':
        return '2025-2026'
    if value == 'before':
        return 'Trước đây'
    return value or ''


def _subject_group_paragraph(admission):
    sg = admission.subject_group
    shift = admission.shift.name if admission.shift_id else ''
    desc = (sg.description or '').strip()
    code = sg.code
    desc_part = f' ({desc})' if desc else ''
    shift_part = f', ca {shift}' if shift else ''
    return (
        'Sau khi tìm hiểu tổ hợp môn chương trình GDTX của trung tâm năm học 2026-2027, '
        f'phụ huynh và học viên chọn Tổ hợp {code}{desc_part}{shift_part} '
        f'để học hết năm đến lớp 12.'
    )


def _campus_address(admission):
    campus = admission.campus
    if campus.address:
        return campus.address.strip()
    return campus.name or ''


def _replace_in_runs(paragraph, mapping):
    for run in paragraph.runs:
        text = run.text
        for key, value in mapping.items():
            text = text.replace(key, value or '')
        run.text = text


def _replace_placeholders(doc, mapping):
    for paragraph in doc.paragraphs:
        _replace_in_runs(paragraph, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph, mapping)


def _build_placeholder_map(admission):
    subjects = [
        s.strip()
        for s in (admission.subject_group.subjects or '').split(',')
        if s.strip()
    ]
    subject_cells = {
        f'{{{{SUBJECT_{idx}}}}}': subjects[idx - 1] if idx - 1 < len(subjects) else ''
        for idx in range(1, 8)
    }
    signed_at = admission.created_at.date() if admission.created_at else None

    return {
        '{{CAMPUS_ADDRESS}}': _campus_address(admission),
        '{{FULL_NAME}}': admission.full_name or '',
        '{{PHONE}}': admission.phone or '',
        '{{GENDER}}': admission.gender or '',
        '{{BIRTHDAY}}': _fmt_date(admission.birthday),
        '{{BIRTH_PLACE}}': admission.birth_place_facility or '',
        '{{ETHNICITY}}': admission.ethnicity or '',
        '{{ID_NUMBER}}': admission.id_number or '',
        '{{ID_ISSUED_DATE}}': _fmt_date(admission.id_issued_date),
        '{{PERM_HOUSE}}': admission.cccd_town or '',
        '{{PERM_STREET}}': admission.cccd_district or '',
        '{{PERM_WARD}}': admission.cccd_ward or '',
        '{{PERM_PROVINCE}}': admission.cccd_province or '',
        '{{CUR_HOUSE}}': admission.birth_reg_town or '',
        '{{CUR_STREET}}': admission.current_district or '',
        '{{CUR_WARD}}': admission.current_ward or '',
        '{{CUR_PROVINCE}}': admission.current_province or '',
        '{{GRADUATION_YEAR}}': _graduation_year_label(admission.graduation_year),
        '{{GRADUATION_SCHOOL}}': admission.graduation_school or '',
        '{{FATHER_NAME}}': admission.father_name or '',
        '{{FATHER_BIRTH}}': admission.father_birth or '',
        '{{FATHER_JOB}}': admission.father_job or '',
        '{{FATHER_PHONE}}': admission.father_phone or '',
        '{{MOTHER_NAME}}': admission.mother_name or '',
        '{{MOTHER_BIRTH}}': admission.mother_birth or '',
        '{{MOTHER_JOB}}': admission.mother_job or '',
        '{{MOTHER_PHONE}}': admission.mother_phone or '',
        '{{SUBJECT_GROUP_PARAGRAPH}}': _subject_group_paragraph(admission),
        '{{SIGN_DATE}}': _fmt_sign_date(signed_at),
        '{{SUBJECT_CODE}}': admission.subject_group.code,
        **subject_cells,
    }


def build_admission_docx(admission):
    if not os.path.isfile(ADMISSION_DOCX_TEMPLATE):
        raise FileNotFoundError('Không tìm thấy file mẫu đơn xin nhập học.')

    doc = Document(ADMISSION_DOCX_TEMPLATE)
    _replace_placeholders(doc, _build_placeholder_map(admission))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def admission_docx_filename(admission):
    name = (admission.full_name or 'ho_so').strip().replace(' ', '_')
    name = re.sub(r'[\\/:*?"<>|]', '', name)
    cccd = re.sub(r'\D', '', admission.id_number or 'CCCD')
    return f'{name}_{cccd}.docx'
